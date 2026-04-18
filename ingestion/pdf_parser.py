"""
Extract clean text from PDF and DOCX files.
- PDFs: PyMuPDF (fitz)
- DOCX: python-docx
"""

import io
import re
from typing import BinaryIO, Union

import fitz
from docx import Document as DocxDocument


def _clean_text(raw: str) -> str:
    """
    Collapse whitespace and strip noise from extracted PDF text.

    Args:
        raw: Raw string from PDF extraction.

    Returns:
        Normalized single string.
    """
    if not raw:
        return ""
    text = raw.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def extract_text_from_pdf(file_obj: BinaryIO) -> str:
    """
    Read a PDF from a binary file-like object and return cleaned text.

    Args:
        file_obj: Open binary stream positioned at start (e.g. uploaded file).

    Returns:
        Full document text, or empty string if unreadable or empty.
    """
    try:
        data = file_obj.read()
        if not data:
            return ""
        doc = fitz.open(stream=data, filetype="pdf")
    except Exception:
        return ""

    parts: list[str] = []
    try:
        for page in doc:
            try:
                parts.append(page.get_text("text") or "")
            except Exception:
                continue
    finally:
        doc.close()

    return _clean_text("\n".join(parts))


def extract_text_from_upload(uploaded_file: Union[BinaryIO, None]) -> str:
    """
    Convenience wrapper for Streamlit UploadedFile objects.

    Args:
        uploaded_file: Streamlit upload or None.

    Returns:
        Cleaned text, or empty string if None or on failure.
    """
    if uploaded_file is None:
        return ""
    try:
        uploaded_file.seek(0)
    except Exception:
        pass
    return extract_text_from_pdf(uploaded_file)


def extract_text_from_bytes(pdf_bytes: bytes) -> str:
    """
    Extract text from raw PDF bytes.

    Args:
        pdf_bytes: Raw PDF content.

    Returns:
        Cleaned text.
    """
    if not pdf_bytes:
        return ""
    return extract_text_from_pdf(io.BytesIO(pdf_bytes))


def extract_text_from_docx(file_obj: BinaryIO) -> str:
    """
    Read a DOCX file from a binary file-like object and return cleaned text.

    Args:
        file_obj: Open binary stream positioned at start (e.g. uploaded file).

    Returns:
        Full document text, or empty string if unreadable or empty.
    """
    try:
        file_obj.seek(0)
        doc = DocxDocument(file_obj)
    except Exception:
        return ""

    parts: list[str] = []
    try:
        for paragraph in doc.paragraphs:
            text = (paragraph.text or "").strip()
            if text:
                parts.append(text)
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells
                )
                if row_text.strip():
                    parts.append(row_text)
    except Exception:
        pass

    return _clean_text("\n".join(parts))


def extract_text_from_docx_upload(uploaded_file: Union[BinaryIO, None]) -> str:
    """
    Convenience wrapper for Streamlit UploadedFile objects (DOCX).

    Args:
        uploaded_file: Streamlit upload or None.

    Returns:
        Cleaned text, or empty string if None or on failure.
    """
    if uploaded_file is None:
        return ""
    try:
        uploaded_file.seek(0)
    except Exception:
        pass
    return extract_text_from_docx(uploaded_file)
